import docx
import pyttsx3
import speech_recognition as sr
import datetime
import wikipedia
import webbrowser
import os
from dotenv import load_dotenv  # Add this line
import smtplib
import google.generativeai as genai

# Initialize Text-to-Speech
engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)

# --- NEW SECURE ENVIRONMENT LOADING ---
# Load environment variables from the .env file
load_dotenv()

# Fetch the API key safely
gemini_api_key = os.getenv("GOOGLE_API_KEY")

# Fail-safe to ensure the key was found
if not gemini_api_key:
    raise ValueError("ERROR: GOOGLE_API_KEY not found. Please check your .env file.")

# Configure Google Gemini API with the secured key
genai.configure(api_key=gemini_api_key)
# --------------------------------------

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 0,
    "max_output_tokens": 8192,
}

safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

model = genai.GenerativeModel(model_name="gemini-1.5-pro-latest", 
                              generation_config=generation_config,
                              safety_settings=safety_settings)

convo = model.start_chat(history=[])

def get_response_from_gemini(user_input):
    """Get a response from Gemini based on the user input."""
    convo.send_message(user_input)
    gemini_reply = convo.last.text
    print(f"Gemini: {gemini_reply}")
    return gemini_reply

# Functions
def speak(audio):
    """Speak the given audio."""
    engine.say(audio)
    engine.runAndWait()

def wishMe():
    """Greet the user based on the time of the day."""
    hour = int(datetime.datetime.now().hour)
    if hour >= 0 and hour < 12:
        speak("Good Morning!")
    elif hour >= 12 and hour < 18:
        speak("Good Afternoon!")   
    else:
        speak("Good Evening!")  
    speak("Hello User. Please tell me how may I help you.")       

def takeCommand():
    """Listen to the user and return their command."""
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        r.pause_threshold = 1
        audio = r.listen(source)
    try:
        print("Recognizing...")    
        query = r.recognize_google(audio, language='en-in')
        print(f"User said: {query}\n")
    except Exception as e:    
        print("Say that again please...")  
        return "None"
    return query

def sendEmail(to, content):
    """Send an email."""
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.ehlo()
        server.starttls()
        server.login('your_email@gmail.com', 'your_password')  # Replace with your email credentials
        server.sendmail('your_email@gmail.com', to, content)
        server.close()
        speak("Email has been sent!")
    except Exception as e:
        print(e)
        speak("Sorry, I couldn't send the email.")

def saveToFile(filename, content):
    """Save content to a specified file."""
    try:
        with open(filename, 'a') as file:
            file.write(content + '\n')
        speak(f"Content has been saved to {filename}.")
    except Exception as e:
        print(e)
        speak("Sorry, I couldn't save the content.")

def readFromFile(filename):
    """Read and return content from a specified file."""
    try:
        with open(filename, 'r') as file:
            content = file.read()
        speak("The file has been read. Would you like me to read the content aloud?")
        response = takeCommand().lower()
        if 'yes' or 'read' or 'aloud' in response:
            speak(content)
        elif 'no' or 'skip':
            speak("Okay, I won't read it aloud.")
        else: 
            speak("Sorry, I did not understand that!")
        print(content)
        return content
    except FileNotFoundError:
        speak("File not found.")
    except Exception as e:
        print(e)
        speak("Sorry, I couldn't read the file.")

# docx functions
def create_docx(filename, content):
  """Create a new DOCX file with the given content."""
  try:
    document = docx.Document()
    paragraph = document.add_paragraph(content)
    document.save(filename + ".docx")
    speak(f"Created a new DOCX file named {filename}.")
  except Exception as e:
    print(e)
    speak("Sorry, I couldn't create the file.")

def read_docx(filename):
  """Read content from a DOCX file and speak it aloud."""
  try:
    document = docx.Document(filename + ".docx")
    full_text = []
    for paragraph in document.paragraphs:
      full_text.append(paragraph.text)
    text = "\n".join(full_text)
    speak(f"The content of {filename} is:")
    speak(text)
  except FileNotFoundError:
    speak("File not found.")
  except Exception as e:
    print(e)
    speak("Sorry, I couldn't read the file.")

def update_docx(filename, content):
  """Update the content of an existing DOCX file."""
  try:
    document = docx.Document(filename + ".docx")
    for paragraph in document.paragraphs:
      paragraph.text = content
      break  # Update only the first paragraph for simplicity
    document.save(filename + ".docx")
    speak(f"Updated the content of {filename}.")
  except FileNotFoundError:
    speak("File not found.")
  except Exception as e:
    print(e)
    speak("Sorry, I couldn't update the file.")

def delete_docx(filename):
  """Delete a DOCX file."""
  try:
    os.remove(filename + ".docx")
    speak(f"Deleted the file {filename}.")
  except FileNotFoundError:
    speak("File not found.")
  except Exception as e:
    print(e)
    speak("Sorry, I couldn't delete the file.")


# Main Functionality
if __name__ == "__main__":
    wishMe()
    listening = True
    sending_to_gemini = False

    while listening:    
        query = takeCommand().lower()

        if 'wikipedia' in query:
            speak('Searching Wikipedia...')
            query = query.replace("wikipedia", "")
            results = wikipedia.summary(query, sentences=2)
            speak("According to Wikipedia")
            print(results)
            speak(results)

        elif 'open youtube' in query:
            webbrowser.open("youtube.com")
            speak("Opening YouTube...")

        elif 'open google' in query:
            webbrowser.open("google.com")
            speak("Opening Google...")
            
        elif 'open stack overflow' in query:
            webbrowser.open("stackoverflow.com")
            speak("Opening Stack Overflow...") 

        elif 'open linkedin' in query:
            webbrowser.open("in.linkedin.com/")
            speak("Opening Linked In...") 
            
        elif 'open facebook' in query:
            webbrowser.open("facebook.com/")
            speak("Opening facebook...")
         
        elif 'open myntra' in query:
            webbrowser.open("myntra.com/")
            speak("Opening myntra...")
        
        elif 'open prime video' in query:
            webbrowser.open("primevideo.com/offers/nonprimehomepage/ref=dv_web_force_root ")
            speak("Opening prime video...")
        
        elif 'open netflix' in query:
            webbrowser.open("netflix.com/in/ ")
            speak("Opening netflix...")
        
        elif 'open cricbuzz' in query:
            webbrowser.open("cricbuzz.com/")
            speak("Opening cricbuzz...")
        
        elif 'open star sports' in query:
            webbrowser.open("starsports.bet/")
            speak("Opening star sports...")  
            
        elif 'prep insta' in query:
            webbrowser.open("prepinsta.com//")
            speak("Opening prep insta...")    
         
        elif 'gfg' in query:
            webbrowser.open("geeksforgeeks.org/")
            speak("Opening prep gfg...") 
            
        
        elif 'play music' in query:
            speak("On what platform should I play them? YouTube Music, Gaana, or Spotify?")
            platform = takeCommand().lower()
            if 'youtube' in platform:
                webbrowser.open("music.youtube.com")
            elif 'gaana' in platform:
                webbrowser.open("gaana.com")
            elif 'spotify' in platform:
                webbrowser.open("open.spotify.com")
            

        elif 'time' in query:
            strTime = datetime.datetime.now().strftime("%H:%M")    
            speak(f"The time is {strTime}")

        elif 'open code' in query:
            codePath = "C:\\Users\\Honeesh\\AppData\\Local\\Programs\\Microsoft VS Code\\Code.exe"  
            os.startfile(codePath)

        elif 'open python' in query:
            codePath = "C:\\Users\\Honeesh\\AppData\\Local\\Programs\\Python\\Python312\\python.exe"  
            os.startfile(codePath)

        elif 'email' in query:
            try:
                speak("What should I say?")
                content = takeCommand()
                to = "kingrhishi@gmail.com"    
                sendEmail(to, content)
            except Exception as e:
                print(e)
                speak("Sorry, I couldn't send this email.") 

        elif 'save to file' in query:
            speak("What should I save?")
            content = takeCommand()
            speak("Would you like to save this to an existing file or create a new file?")
            choice = takeCommand().lower()

            if 'existing' in choice:
                speak("What is the name of the existing file?")
                filename = takeCommand().lower().replace(" ", "_") + ".txt"
                saveToFile(filename, content)
            elif 'new' in choice:
                speak("What should be the name of the new file?")
                filename = takeCommand().lower().replace(" ", "_") + ".txt"
                saveToFile(filename, content)
            else:
                speak("Sorry, I did not understand that. Please try again.")

        elif 'read file' in query:
            speak("Which file should I read?")
            filename = takeCommand().lower().replace(" ", "_") + ".txt"
            readFromFile(filename)

        elif 'bye' in query or 'quit' in query:
            speak("Goodbye, have a great day!")
            break
   
        elif 'create' in query and 'document' in query:
            speak("What should be the filename?")
            filename = takeCommand().lower().replace(" ", "_")
            speak("What content should I write?")
            content = takeCommand()
            create_docx(filename, content)

        elif 'read' in query and 'document' in query:
            speak("What file should I read?")
            filename = takeCommand().lower().replace(" ", "_")
            read_docx(filename)

        elif 'update' in query and 'document' in query:
            speak("What file should I update?")
            filename = takeCommand().lower().replace(" ", "_")
            speak("What should the new content be?")
            content = takeCommand()
            update_docx(filename, content)

        elif 'delete' in query and 'document' in query:
            speak("What file should I delete?")
            filename = takeCommand().lower().replace(" ", "_")
            delete_docx(filename)

        elif 'gemini' in query:
            speak("What would you like to ask Gemini?")
            user_input = takeCommand()
            if user_input.lower() not in ['exit', 'stop']:
                gemini_reply = get_response_from_gemini(user_input)
                speak(gemini_reply)
            else:
                speak("Stopping Gemini interactions.")
                sending_to_gemini = False

        elif 'bye' in query or 'quit' in query:
            speak("Goodbye, have a great day!")
            break
